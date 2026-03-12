from docx import Document
import re
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def parse_docx(file_path):
    doc = Document(file_path)
    questions = []
    current_q = None
    current_option = None
    option_buffer = ""

    def save_option():
        nonlocal current_q, current_option, option_buffer
        if current_q and current_option:
            current_q["options"][current_option] = option_buffer.strip()
        option_buffer = ""
        current_option = None

    def is_correct_format(run, paragraph):

        if run.bold:
            return True

        if run.underline:
            return True

        if run.style and run.style.font and run.style.font.bold:
            return True

        if paragraph.style and paragraph.style.font and paragraph.style.font.bold:
            return True

        if run.font and run.font.highlight_color:
            return True

        if run.font and run.font.color and run.font.color.rgb:
            return True

        return False

    def handle_paragraph(paragraph):
        nonlocal current_q, questions, current_option, option_buffer

        text = paragraph.text.strip()
        if not text:
            return

        lines = text.split("\n")

        for text in lines:
            text = text.strip()
            if not text:
                continue

            question_match = re.match(r'^(câu\s*\d+[\.\):]?|\d+[\.\):])\s*(.*)', text, re.IGNORECASE)

            if question_match:
                save_option()

                if current_q and current_q["options"]:
                    questions.append(current_q)

                current_q = {
                    "text": text,
                    "options": {},
                    "correct": None
                }
                continue

            multi_option = re.findall(r'([A-D])[\.\):]\s*(.*?)(?=\s+[A-D][\.\):]|$)', text)

            if len(multi_option) >= 2 and current_q:
                save_option()

                for letter, content in multi_option:
                    current_q["options"][letter] = content.strip()

                    for run in paragraph.runs:
                        if letter in run.text and is_correct_format(run, paragraph):
                            current_q["correct"] = letter

                return

            option_match = re.match(r'^([A-D])[\.\):]\s*(.*)', text)

            if not option_match and current_q and len(current_q["options"]) < 4:
                option_letter = ["A", "B", "C", "D"][len(current_q["options"])]
                current_q["options"][option_letter] = text

                for run in paragraph.runs:
                    if is_correct_format(run, paragraph) and run.text.strip():
                        current_q["correct"] = option_letter
                        break

                continue

            if option_match:
                if option_match.group(1).upper() == 'A' and current_q and current_q.get("text"):
                    save_option()
                    if current_q["options"]:
                        questions.append(current_q)

                save_option()
                current_option = option_match.group(1).upper()
                option_buffer = option_match.group(2)

                for run in paragraph.runs:
                    if is_correct_format(run, paragraph):
                        current_q["correct"] = current_option
                        break

                continue

            if current_q is None or (current_q and current_q["options"]):
                if current_q and current_q["options"]:
                    save_option()
                    questions.append(current_q)
                    current_q = {"text": text, "options": {}, "correct": None}
                else:
                    if current_q is None:
                        current_q = {"text": text, "options": {}, "correct": None}
                    else:
                        current_q["text"] += " " + text
            else:
                current_q["text"] += " " + text

    for para in doc.paragraphs:
        handle_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    handle_paragraph(para)

    if current_q:
        save_option()
        if not current_q["correct"]:
            current_q["correct"] = "incorrect"
        questions.append(current_q)

        for i, q in enumerate(questions, start=1):
            if not re.match(r'^(câu\s*\d+|\d+[\.\)])', q["text"], re.IGNORECASE):
                q["text"] = f"Câu {i}. " + q["text"]

    return questions


@app.post("/upload-exam")
async def upload_exam(file: UploadFile = File(...)):
    file_path = f"temp_{file.filename}"

    with open(file_path, "wb") as f:
        f.write(await file.read())

    questions = parse_docx(file_path)

    os.remove(file_path)

    return {
        "message": "Upload & parse thành công",
        "total_questions": len(questions),
        "questions": questions
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)